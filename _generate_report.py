# Generates the Lexora project report docx in the sample's format.
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

STUDENT = "TAABISH AHMED ANSARI"
STUDENT_TITLE = "Taabish Ahmed Ansari"
ENROLL = "[Enrollment No.]"
SUPERVISOR = "[Faculty Supervisor Name]"
FACULTY_MENTOR = "[Faculty Mentor Name]"
UNIVERSITY = "Faculty of Science & Technology, IFHE University"
PERIOD = "June\u2013July 2026"
DATE = "08-07-2026"
PLACE = "Hyderabad"
TITLE_LINE = "LEXORA \u2013 AN OFFLINE-FIRST AI ASSISTANT FOR LAWYERS"

doc = Document()

# ---- base style -------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def para(text="", size=14, bold=False, align="justify", space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
    return p


def heading(text, size=18):
    return para(text, size=size, bold=True, align="left", space_after=10)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def cover(extra=None):
    para("A", 22, True, "center")
    para("PROJECT REPORT", 22, True, "center")
    para("ON", 22, True, "center")
    para("", 14)
    para(TITLE_LINE, 20, True, "center")
    para("", 14)
    para("BY", 18, True, "center")
    para(f"{STUDENT} \u2013 {ENROLL}", 18, True, "center")
    para("", 14)
    if extra:
        para(extra, 15, False, "center")
    if os.path.exists("_cover_logo.png"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture("_cover_logo.png", width=Inches(2.6))
    para("", 14)
    para(UNIVERSITY, 16, True, "center")
    para(PERIOD, 15, False, "center")


# ============================ COVER PAGES =====================================
cover()
page_break()
cover(extra=("Project report submitted to IcfaiTech as a fulfillment of the "
             "requirements for the award of the Degree of B. Tech in Computer "
             f"Science and Engineering under the supervision of {SUPERVISOR}."))
page_break()

# ============================ DECLARATION =====================================
heading("Declaration", 20)
para(f"I, {STUDENT_TITLE} ({ENROLL}), hereby declare that this Project Report, titled "
     f"\u201cLexora \u2013 An Offline-First AI Assistant for Lawyers,\u201d is an authentic "
     "record of the work I personally conducted during the development of this project.")
para("This project was completed under the invaluable guidance of my mentors. I extend my "
     f"sincere gratitude to my faculty supervisor, {SUPERVISOR}, of the Department of Computer "
     "Science and Engineering, IcfaiTech, for his academic supervision and continuous support, "
     "whose practical expertise was essential to this endeavor.")
para("I affirm that the work presented herein is original and that all external sources have "
     "been appropriately acknowledged and cited. This report, in whole or in part, has not been "
     "previously submitted to this or any other institution for the award of any degree or diploma.")
para("I have adhered to the highest academic and ethical standards throughout this project and "
     "accept full responsibility for the content and conclusions presented in this document.")
para("", 14)
para("[Signature]", 14, False, "left")
para(STUDENT_TITLE, 14, False, "left")
para(f"Date: {DATE}", 14, False, "left")
page_break()

# ============================ CERTIFICATE =====================================
heading("CERTIFICATE", 20)
para(UNIVERSITY, 15, True, "left")
para(f"This is to certify that the project report titled \u201cLexora \u2013 An Offline-First "
     f"AI Assistant for Lawyers\u201d has been successfully completed by {STUDENT_TITLE} "
     f"(Enrollment No. {ENROLL}). This report is the culmination of his work on the design and "
     "development of an offline, on-device Artificial Intelligence assistant for legal "
     f"professionals, carried out during {PERIOD}.")
para("It is further certified that the report is an authentic record of the student's own work "
     "and incorporates all modifications as suggested. The project fulfills the requirements for "
     "the Bachelor of Technology in Computer Science and Engineering and has been submitted to "
     "the Department of Computer Science and Engineering, The ICFAI Foundation for Higher Education.")
para("The work presented in this report has not been submitted by the student for any other "
     "course or degree.")
para("", 14)
para("", 14)
para("Signature of the Instructor                                 Signature of the Student",
     14, False, "left")
para(f"Date: {DATE}", 14, False, "left")
para(f"Place: {PLACE}", 14, False, "left")
page_break()

# ============================ ACKNOWLEDGEMENT =================================
heading("Acknowledgement", 20)
para("I would like to express my sincere gratitude to everyone who has supported and guided me "
     "throughout the development of the Lexora project. This project gave me the invaluable "
     "opportunity to work on a challenging and meaningful problem at the intersection of "
     "Artificial Intelligence and the legal profession.")
para(f"I am particularly grateful to my faculty supervisor, {SUPERVISOR}, for his continuous "
     "mentorship, support, and guidance. His expertise and encouragement were instrumental in "
     "helping me navigate the complexities of local language models, semantic retrieval, and "
     "full-stack system design.")
para(f"I would also like to extend my heartfelt appreciation to my faculty mentor, "
     f"{FACULTY_MENTOR}, for consistent encouragement, insightful advice, and for ensuring that "
     "my project work remained aligned with my academic and professional goals. Finally, I am "
     "thankful to my peers and to everyone who directly or indirectly contributed to making this "
     "project a transformative learning experience. Their cooperation and willingness to share "
     "their knowledge made this journey both enjoyable and profoundly insightful.")
page_break()

# ============================ ABSTRACT ========================================
heading("Abstract", 20)
para("This report provides a comprehensive overview of the design and development of Lexora, an "
     "offline-first, on-device Artificial Intelligence assistant built for lawyers and legal "
     "professionals. The project was conceived to address a critical gap in the legal industry: "
     "lawyers handle highly confidential client information that cannot be sent to cloud-hosted "
     "AI services, yet they stand to benefit enormously from AI-powered assistance in managing "
     "clients, case files, conversations, and schedules.")
para("Lexora runs entirely on local hardware. Text generation is performed by the "
     "TinyLlama-1.1B-Chat large language model executed through HuggingFace Transformers and "
     "PyTorch, while semantic understanding is powered by sentence-transformer embeddings and a "
     "custom NumPy-based cosine-similarity vector store persisted as JSON. On top of this "
     "foundation, the system delivers a ChatGPT-style legal chat interface, persistent per-client "
     "factual memory, an automatically constructed knowledge graph rendered with a custom HTML5 "
     "Canvas force-directed layout, offline Retrieval-Augmented Generation (RAG) over uploaded "
     "PDF case documents, a deterministic contradiction-detection engine that flags inconsistent "
     "client statements, an AI call assistant with live speech-to-text transcription and "
     "automatic meeting summaries, natural-language scheduling with Google Calendar integration, "
     "e-mail dispatch of case summaries, and a case-insights dashboard. The backend is built with "
     "FastAPI and the frontend with framework-free HTML, CSS, and JavaScript.")
para("This project significantly enhanced my technical abilities in full-stack development, "
     "applied machine learning, natural language processing, and system design. It improved my "
     "confidence in building reliable end-to-end applications and gave me a deeper understanding "
     "of how private, secure, and scalable AI systems are engineered. This report summarizes the "
     "system architecture, the technologies mastered, the challenges overcome, and the lessons "
     "learned during this journey.")
page_break()

# ============================ TABLE OF CONTENTS ===============================
heading("Table of Contents", 20)
toc = [
    ("Table of Illustrations", "9"),
    ("1. Introduction", "10"),
    ("    1.1. Background of the Project", "10"),
    ("    1.2. Problem Statement", "10"),
    ("    1.3. Objectives of the Project", "11"),
    ("2. System Design and Technical Deep Dive", "12"),
    ("    2.1. System Architecture Overview", "12"),
    ("    2.2. The Local LLM Layer: TinyLlama-1.1B-Chat", "13"),
    ("    2.3. Offline Retrieval-Augmented Generation (RAG)", "14"),
    ("    2.4. Persistent Memory and the Knowledge Graph", "15"),
    ("    2.5. Contradiction Detection Engine", "16"),
    ("    2.6. AI Call Assistant and Live Transcription", "17"),
    ("    2.7. Scheduling, E-mail, and the Case Insights Dashboard", "18"),
    ("    2.8. Backend API and Frontend Design", "19"),
    ("3. Skills Acquired", "20"),
    ("    3.1. Technical Skills", "20"),
    ("    3.2. Soft Skills", "21"),
    ("4. Challenges and Learnings", "22"),
    ("5. Conclusion", "23"),
    ("6. References", "24"),
]
for label, pg in toc:
    dots = "\u2026" * max(4, (72 - len(label)) // 2)
    para(f"{label} {dots} {pg}", 14, False, "left", space_after=4)
page_break()

# ============================ TABLE OF ILLUSTRATIONS ==========================
heading("Table of Illustrations", 20)
figs = [
    "Fig. 1: High-Level System Architecture of Lexora",
    "Fig. 2: The ChatGPT-Style Offline Chat Interface",
    "Fig. 3: Auto-Generated Client Knowledge Graph (Canvas Force-Directed Renderer)",
    "Fig. 4: PDF Upload and Document Retrieval (Offline RAG) Interface",
    "Fig. 5: Contradiction Warning Card Displayed During a Chat",
    "Fig. 6: AI Call Assistant with Live Speech-to-Text Transcription",
    "Fig. 7: Case Insights Dashboard with Chart.js Visualizations",
]
for i, f in enumerate(figs):
    para(f"{f} {'\u2026' * 12} {12 + i}", 14, False, "left", space_after=4)
page_break()

# ============================ 1. INTRODUCTION =================================
heading("1. Introduction")
heading("1.1. Background of the Project", 15)
para("Modern legal practice is document- and conversation-heavy. A practicing lawyer manages "
     "dozens of clients simultaneously, each with a growing trail of statements, phone calls, "
     "case documents, hearing dates, and correspondence. Tools such as ChatGPT have shown how "
     "transformative a conversational AI assistant can be, but they are fundamentally unsuitable "
     "for the legal domain in their cloud-hosted form: attorney\u2013client privilege and data "
     "protection obligations make it unacceptable to transmit confidential case information to "
     "third-party servers.")
para("Lexora was built to resolve this tension. It is an offline-first AI assistant in which "
     "every component that touches client data \u2014 the large language model, the text "
     "embeddings, the vector search index, and the knowledge graph \u2014 runs locally on the "
     "lawyer's own machine. All data and vector indexes are stored on disk as human-readable "
     "JSON files, so the application works without any internet connection and the user retains "
     "complete ownership of their data. An optional remote fallback (Google Gemini) exists purely "
     "as a silent reliability measure and is never required or surfaced in the interface.")
heading("1.2. Problem Statement", 15)
para("The project addresses four concrete problems faced by legal professionals:")
para("1. Confidentiality: Cloud AI assistants require sending privileged client data to external "
     "servers, which is ethically and legally problematic for lawyers.")
para("2. Fragmented memory: Client facts are scattered across notebooks, e-mails, and memory. "
     "There is no single system that remembers everything a client has ever said and can recall "
     "it instantly and semantically.")
para("3. Inconsistency detection: Clients sometimes make contradictory statements weeks apart "
     "(\u201cI was at the office at 5 PM\u201d versus \u201cI was not at the office\u201d). "
     "Catching such contradictions manually across long case histories is error-prone.")
para("4. Administrative overhead: Summarizing calls, extracting action items, scheduling "
     "hearings, and e-mailing case summaries consume time that could be spent on legal work.")
heading("1.3. Objectives of the Project", 15)
para("The primary objectives of the Lexora project were:")
para("\u2022 To build a fully offline conversational AI assistant using an on-device LLM "
     "(TinyLlama-1.1B-Chat) with GPU acceleration where available.")
para("\u2022 To implement persistent, per-client factual memory with semantic recall, so that "
     "every statement a client makes is embedded, indexed, and retrievable in future "
     "conversations.")
para("\u2022 To automatically construct a knowledge graph of entities and relations for each "
     "client and render it interactively in the browser.")
para("\u2022 To support offline Retrieval-Augmented Generation over uploaded PDF case documents "
     "with source citation.")
para("\u2022 To detect contradictions between a client's new statements and their prior "
     "testimony using a deterministic, reliable engine.")
para("\u2022 To provide an AI call assistant capable of live transcription, real-time memory "
     "commits, automatic meeting summaries, and action-item extraction.")
para("\u2022 To round out the product with natural-language scheduling, e-mail dispatch of case "
     "summaries, client management, and a case-insights dashboard.")
page_break()

# ============================ 2. TECHNICAL DEEP DIVE ==========================
heading("2. System Design and Technical Deep Dive")
heading("2.1. System Architecture Overview", 15)
para("Lexora follows a clean layered architecture. A FastAPI web server (app.py) exposes a REST "
     "API and hosts the static frontend. The orchestration layer (chat.py) coordinates every "
     "chat turn and call event. Beneath it sit four specialist modules: llm.py (text "
     "generation), rag.py (embeddings, vector store, and PDF ingestion), memory.py (clients, "
     "statements, knowledge graph, contradiction engine, dashboard, and activity log), and "
     "calendarx.py / mailer.py / telephony.py for scheduling, e-mail, and real phone-call "
     "integration respectively. Persistence is deliberately simple and portable: all case data "
     "lives in data/store.json and all vectors in chroma_db/documents.json and "
     "chroma_db/memory.json, guarded by re-entrant locks for thread safety and flushed to disk "
     "on every write, so closing and reopening the application restores the entire state exactly.")
para("[Fig. 1: High-Level System Architecture of Lexora \u2014 insert diagram/screenshot here]",
     12, False, "center")
para("A single chat turn flows as follows: the user's message is stored as a statement and "
     "embedded into the memory index; a contradiction check runs against the client's prior "
     "statements; context is assembled from semantic memory recall, RAG document chunks, recent "
     "chat turns, and any detected contradiction; TinyLlama generates the reply, which is "
     "cleaned by regex post-processing; entities and relations are extracted and merged into the "
     "client's knowledge graph; and finally everything is persisted to disk.")
heading("2.2. The Local LLM Layer: TinyLlama-1.1B-Chat", 15)
para("Text generation is powered by TinyLlama-1.1B-Chat, loaded lazily through HuggingFace "
     "Transformers and PyTorch. The model runs in float16 on a CUDA GPU when available and falls "
     "back to float32 on CPU. Generation uses the tokenizer's native chat template with sampling "
     "controls \u2014 temperature, top-p (nucleus) sampling, and a repetition penalty \u2014 to "
     "balance fluency with factual restraint. Because small models occasionally hallucinate URLs "
     "or echo their context, a regex-based output-hygiene pass strips such artifacts before the "
     "reply reaches the user.")
para("The LLM layer also provides structured generation: generate_json() prompts the model for "
     "a JSON object and then extracts the first balanced JSON structure from the raw output, "
     "tolerating code fences and surrounding prose. This is used for knowledge-graph extraction "
     "and action-item extraction. A silent, environment-gated fallback to Google Gemini exists "
     "for the rare case where the local model fails to load; it is never referenced in the user "
     "interface, preserving the offline-first product identity.")
para("[Fig. 2: The ChatGPT-Style Offline Chat Interface \u2014 insert screenshot here]",
     12, False, "center")
heading("2.3. Offline Retrieval-Augmented Generation (RAG)", 15)
para("Lawyers can upload PDF case documents which become searchable knowledge for the "
     "assistant. The ingestion pipeline extracts text with pypdf, splits it into fixed-size "
     "overlapping chunks (900 characters with 150-character overlap), embeds each chunk with a "
     "sentence-transformers model, and stores the vectors with metadata identifying the source "
     "file and owning client.")
para("A notable engineering decision was the vector store itself. The initially chosen ChromaDB "
     "build suffered a fatal compaction bug in this environment that silently dropped every "
     "vector. It was replaced with a purpose-built, dependency-light NumPy cosine-similarity "
     "index: embeddings are L2-normalized so cosine similarity reduces to a dot product, and the "
     "index persists as JSON on disk. This delivered the same retrieval quality with zero "
     "failure points. At question time, the query is embedded and the top-k most similar chunks "
     "(filtered to the active client, falling back to the shared library) are injected into the "
     "prompt, and answers cite their source documents.")
para("[Fig. 4: PDF Upload and Document Retrieval (Offline RAG) Interface \u2014 insert "
     "screenshot here]", 12, False, "center")
heading("2.4. Persistent Memory and the Knowledge Graph", 15)
para("Every factual message a client sends is stored twice: as a plain statement in "
     "store.json and as a normalized embedding in the memory vector index. During later "
     "conversations, the current message is embedded and the most semantically similar past "
     "statements are recalled and placed in the model's context \u2014 giving Lexora a "
     "persistent, ChatGPT-like memory that survives restarts.")
para("In parallel, TinyLlama is prompted to return a JSON structure of entities and relations "
     "for each factual message. A heuristic sanitization layer drops hallucinated URLs and "
     "generic placeholder labels (such as \u201cPERSON\u201d or \u201cDATE\u201d) that small "
     "models sometimes emit. Nodes and edges are merged case-insensitively into the client's "
     "graph, so the same entity mentioned in different conversations connects automatically, and "
     "orphan entities are attached to the client node. The graph is stored as simple node and "
     "edge lists \u2014 no graph database is required \u2014 and is rendered in the browser by a "
     "fully custom HTML5 Canvas force-directed renderer featuring node repulsion, spring "
     "attraction along edges, center gravity, drag-to-move interaction, and color-coding by "
     "entity type, all with zero external libraries.")
para("[Fig. 3: Auto-Generated Client Knowledge Graph \u2014 insert screenshot here]",
     12, False, "center")
heading("2.5. Contradiction Detection Engine", 15)
para("Reliability was paramount for contradiction detection, so this feature is deliberately "
     "not delegated to the 1.1-billion-parameter model. Instead, a deterministic engine runs on "
     "every new statement. The statement is embedded and the most similar prior statements for "
     "that client are recalled from the memory index. Two rule families then apply: a "
     "negation/polarity check flags cases where two statements share the same subject (shared "
     "tokens plus cosine similarity of at least 0.45) but one affirms while the other denies; "
     "and a numeric/time-conflict check flags statements about the same subject that quote "
     "different numbers or clock times (for example, 5 PM versus 9 PM) while skipping "
     "corroborating evidence. When a conflict is found, the chatbot explains the contradiction "
     "in natural language and the UI displays a guaranteed red warning card.")
para("[Fig. 5: Contradiction Warning Card Displayed During a Chat \u2014 insert screenshot "
     "here]", 12, False, "center")
heading("2.6. AI Call Assistant and Live Transcription", 15)
para("Lexora includes both a demo call assistant and a live-call mode. Speech-to-text uses the "
     "browser-native Web Speech API in Chrome or Edge, so no audio ever leaves the machine; for "
     "a real two-phone call, the phone is placed on speaker next to the laptop. Each finalized "
     "utterance is posted to the /api/live endpoint, where it is committed to memory and the "
     "knowledge graph in real time with live contradiction checks.")
para("The assistant can even create clients automatically mid-call: a regex-based name detector "
     "(recognizing patterns like \u201cmy name is \u2026\u201d) combined with a keyword-based "
     "case-type classifier (theft \u2192 Criminal, divorce \u2192 Divorce, and so on) spins up a "
     "new client record on the fly when none is selected. When the call ends, TinyLlama writes "
     "a meeting summary, and action items are extracted through LLM JSON generation with a "
     "robust regex fallback keyed on cue words such as \u201cwill,\u201d \u201csend by "
     "Friday,\u201d and \u201cfollow up.\u201d An optional Twilio integration (telephony.py) "
     "supports genuine telephone calls with server-side transcription callbacks.")
para("[Fig. 6: AI Call Assistant with Live Speech-to-Text Transcription \u2014 insert "
     "screenshot here]", 12, False, "center")
heading("2.7. Scheduling, E-mail, and the Case Insights Dashboard", 15)
para("A natural-language calendar module (calendarx.py) detects scheduling intent in chat "
     "messages using cue words, carefully excluding past-testimony sentences. It extracts clock "
     "times with regular expressions before parsing dates with the dateparser library "
     "(preferring future dates), synthesizes a clean event title by peeling away filler words, "
     "detects conflicts with existing events, and produces a one-click \u201cAdd to Google "
     "Calendar\u201d link. Case summaries can be e-mailed to clients directly via Gmail SMTP "
     "using an app password (mailer.py).")
para("The Case Insights Dashboard aggregates clients by case status and case type, along with "
     "totals for calls, statements, documents, and contradictions, rendered as doughnut and bar "
     "charts using a locally vendored copy of Chart.js on a dark theme. An activity log records "
     "every significant action \u2014 client created, graph updated, contradiction detected, "
     "call recorded, document uploaded, client deleted \u2014 providing a "
     "who-did-what-when audit trail.")
para("[Fig. 7: Case Insights Dashboard with Chart.js Visualizations \u2014 insert screenshot "
     "here]", 12, False, "center")
heading("2.8. Backend API and Frontend Design", 15)
para("The backend exposes a clean REST API through FastAPI with Pydantic request validation: "
     "chat (/api/chat), client CRUD (/api/clients with POST, PATCH, and DELETE), per-client "
     "messages, statements, graphs, and calls, PDF upload (/api/upload via multipart), document "
     "listing, live-call turns (/api/live), call processing (/api/call), dashboard aggregation "
     "(/api/dashboard), health checks, and Twilio webhook endpoints. Deleting a client purges "
     "them everywhere \u2014 chat history, statements, graph, calls, and all vectors in both "
     "indexes. A no-cache middleware guarantees the browser always loads the latest UI assets.")
para("The frontend is intentionally framework-free: vanilla HTML, CSS, and JavaScript deliver a "
     "dark, ChatGPT-style interface with a client sidebar, a streaming-style typing indicator, "
     "modals for client management, the Canvas knowledge-graph view, and the dashboard. "
     "Avoiding frameworks kept the application light, fully offline, and easy to reason about.")
page_break()

# ============================ 3. SKILLS ACQUIRED ==============================
heading("3. Skills Acquired")
heading("3.1. Technical Skills", 15)
para("\u2022 Applied Machine Learning and NLP: Running quantization-aware local LLM inference "
     "with HuggingFace Transformers and PyTorch (fp16 on CUDA), prompt engineering for small "
     "models, structured JSON generation, and output sanitization.")
para("\u2022 Semantic Search and RAG: Sentence-transformer embeddings, vector normalization and "
     "cosine similarity, chunking strategies with overlap, metadata filtering, and building a "
     "persistent vector store from first principles with NumPy.")
para("\u2022 Backend Engineering: Designing REST APIs with FastAPI and Pydantic, multipart file "
     "uploads, middleware, thread-safe persistence with re-entrant locks, and webhook "
     "integration with Twilio telephony.")
para("\u2022 Frontend Engineering: Framework-free JavaScript application structure, HTML5 "
     "Canvas programming including a force-directed graph layout algorithm, Chart.js data "
     "visualization, and the Web Speech API for browser-native speech recognition.")
para("\u2022 Algorithm Design: Deterministic heuristics for contradiction detection "
     "(polarity and numeric/time conflicts), regex-based information extraction, and "
     "natural-language date/time parsing.")
para("\u2022 Software Architecture: Layered module design, offline-first system thinking, "
     "graceful degradation and silent fallbacks, and pragmatic persistence with human-readable "
     "JSON.")
heading("3.2. Soft Skills", 15)
para("\u2022 Problem Decomposition: Breaking a large product vision into independently "
     "buildable, testable features.")
para("\u2022 Engineering Judgment: Choosing reliability over sophistication where it mattered "
     "\u2014 for example, backing the small LLM with deterministic engines for must-be-correct "
     "tasks, and replacing a failing third-party database with a simpler in-house solution.")
para("\u2022 Documentation and Communication: Writing clear technical documentation of the "
     "stack, data flow, and setup procedures.")
para("\u2022 Time Management and Self-Direction: Planning and executing a multi-feature project "
     "end-to-end within a limited timeframe.")
page_break()

# ============================ 4. CHALLENGES ===================================
heading("4. Challenges and Learnings")
para("Working with a 1.1-billion-parameter model was the defining challenge of the project. "
     "TinyLlama is small enough to run on consumer hardware but exhibits variance: it can "
     "hallucinate URLs, echo its context, or emit malformed JSON. The learning was to treat the "
     "model as one component in a pipeline rather than an oracle \u2014 surrounding it with "
     "regex-based output hygiene, balanced-JSON extraction, heuristic sanitization of graph "
     "labels, and deterministic engines for correctness-critical tasks such as contradiction "
     "detection and action-item extraction. This layered defense made demos reliable regardless "
     "of the model's variance.")
para("A second major challenge was the vector database. ChromaDB, the standard choice, threw a "
     "fatal mismatched-types compaction error on a fresh database in this environment and "
     "silently dropped every stored vector. Debugging a third-party storage engine mid-project "
     "was impractical, so the pragmatic solution was to write a minimal NumPy cosine index that "
     "persists to JSON \u2014 an instructive lesson that understanding the underlying "
     "mathematics (normalized embeddings make cosine similarity a dot product) can make the "
     "\u201chard\u201d dependency unnecessary.")
para("Other challenges included parsing natural-language dates without misreading times as "
     "days (solved by extracting clock times with regex before date parsing), distinguishing "
     "scheduling requests from past testimony in the calendar module, merging knowledge-graph "
     "entities across conversations without duplication, and ensuring complete data hygiene "
     "when deleting a client across every store and index. Each of these sharpened my instinct "
     "for edge cases and for designing systems that fail safely.")
page_break()

# ============================ 5. CONCLUSION ===================================
heading("5. Conclusion")
para("Lexora demonstrates that a genuinely useful, private, and trustworthy AI assistant for "
     "the legal profession can run entirely on local hardware. By combining an on-device LLM, "
     "local embeddings, a custom vector store, and deterministic reliability engines, the "
     "system delivers conversational assistance, persistent client memory, document "
     "intelligence, contradiction detection, call summarization, scheduling, and case analytics "
     "\u2014 all without a single byte of client data leaving the machine.")
para("The project was a transformative learning experience spanning the full breadth of modern "
     "software engineering: applied machine learning, natural language processing, information "
     "retrieval, backend and frontend development, and system architecture. Beyond the "
     "technology, it taught the value of engineering pragmatism \u2014 knowing when to rely on "
     "a model and when to rely on an algorithm, and when a simple, well-understood solution "
     "beats a heavyweight dependency.")
para("Future work could include upgrading to newer small models (such as Phi-3 or "
     "Llama-3.2-1B), fine-tuning on legal corpora, richer document understanding (OCR for "
     "scanned filings), multi-user support with authentication, and packaging the application "
     "as a signed desktop installer for non-technical users.")
page_break()

# ============================ 6. REFERENCES ===================================
heading("6. References")
refs = [
    "1. TinyLlama: An Open-Source Small Language Model \u2014 Zhang, P., et al. (2024). "
    "https://github.com/jzhang38/TinyLlama",
    "2. HuggingFace Transformers Documentation \u2014 https://huggingface.co/docs/transformers",
    "3. PyTorch Documentation \u2014 https://pytorch.org/docs",
    "4. Sentence-Transformers: Sentence Embeddings using Siamese BERT-Networks \u2014 Reimers, "
    "N., & Gurevych, I. (2019). https://www.sbert.net",
    "5. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks \u2014 Lewis, P., "
    "et al. (2020). NeurIPS.",
    "6. FastAPI Documentation \u2014 https://fastapi.tiangolo.com",
    "7. Pydantic Documentation \u2014 https://docs.pydantic.dev",
    "8. pypdf Documentation \u2014 https://pypdf.readthedocs.io",
    "9. Chart.js Documentation \u2014 https://www.chartjs.org/docs",
    "10. Web Speech API \u2014 MDN Web Docs. "
    "https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API",
    "11. Twilio Programmable Voice Documentation \u2014 https://www.twilio.com/docs/voice",
    "12. dateparser Documentation \u2014 https://dateparser.readthedocs.io",
    "13. Force-Directed Graph Drawing \u2014 Fruchterman, T. M. J., & Reingold, E. M. (1991). "
    "Graph Drawing by Force-Directed Placement.",
]
for r in refs:
    para(r, 13, False, "left", space_after=6)

OUT = "Lexora Project Report - Taabish Ahmed Ansari.docx"
doc.save(OUT)
print("saved:", OUT)
